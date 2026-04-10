"""
writer_docx.py — Render a quiz object to a DOCX file using python-docx.

Styles are built programmatically — no reference template required.
"""

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# ── Helpers ──────────────────────────────────────────────────────────────────

def _set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        _set_font(run, size=13 if level == 1 else 11, bold=True)
    return p


def _para(doc, text="", bold=False, italic=False, size=11, left_indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(left_indent)
    p.paragraph_format.space_after = Pt(4)
    if text:
        run = p.add_run(text)
        _set_font(run, size=size, bold=bold, italic=italic)
    return p


def _inline_bold(p, parts):
    """
    Add a sequence of (text, bold) tuples to paragraph p.
    parts: list of (str, bool) — e.g. [("She said ", False), ("despite", True), (" loudly.", False)]
    """
    for text, bold in parts:
        run = p.add_run(text)
        _set_font(run, bold=bold)


def _add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _question_number(sections, section, index):
    order = ["A", "B", "C", "D"]
    offset = 1
    for s in order:
        if s == section:
            return offset + index
        offset += len(sections[s])
    return offset + index


def _parse_bold_markers(text: str) -> list[tuple[str, bool]]:
    """
    Split text on **...** markers into (segment, is_bold) pairs.
    e.g. "She tried to **prevent** her." → [("She tried to ", False), ("prevent", True), (" her.", False)]
    """
    parts = []
    while "**" in text:
        before, rest = text.split("**", 1)
        if "**" in rest:
            bold_text, text = rest.split("**", 1)
            if before:
                parts.append((before, False))
            parts.append((bold_text, True))
        else:
            parts.append((text, False))
            text = ""
            break
    if text:
        parts.append((text, False))
    return parts if parts else [(text, False)]


# ── Main writer ──────────────────────────────────────────────────────────────

def write(quiz: dict, output_dir: Path) -> Path:
    if not DOCX_AVAILABLE:
        raise ImportError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    n = quiz["number"]
    book_label = quiz["book"].upper()
    page_range = quiz["page_range"]
    words = quiz["words"]
    sections = quiz["sections"]

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Header ───────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"Great Writing {book_label[-1]} — Vocabulary Quiz {n:02d}")
    _set_font(r, size=14, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(
        f"Words: {page_range}  |  Total: {len(words)} words  |  Time: 20 minutes"
    )
    _set_font(r2, size=10, italic=True, color=(100, 100, 100))

    doc.add_paragraph()  # spacer

    # Student info row
    info = doc.add_paragraph()
    r3 = info.add_run("Name: ________________________________    ")
    _set_font(r3, size=11)
    r4 = info.add_run("Class: _______    ")
    _set_font(r4, size=11)
    r5 = info.add_run("Date: ___________")
    _set_font(r5, size=11)

    _add_divider(doc)
    doc.add_paragraph()

    # ── Instructions ─────────────────────────────────────────────────────────
    _heading(doc, "Instructions", level=2)
    for text in [
        "Section A: Circle the letter of the best answer.",
        "Section B: Write the missing word on the line. Use the synonym clue to help you.",
        "Section C: Each sentence has one incorrect word (underlined). Write the correct word.",
        "Section D: Write 1–2 sentences using the vocabulary word. Underline it in your answer.",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(text)
        _set_font(run, size=10)

    _add_divider(doc)

    # ── Section A — Multiple Choice ──────────────────────────────────────────
    _heading(doc, "Section A — Multiple Choice", level=2)

    for i, item in enumerate(sections["A"]):
        q = _question_number(sections, "A", i)
        p = _para(doc)
        r_num = p.add_run(f"{q}.  ")
        _set_font(r_num, bold=True)
        r_stem = p.add_run(item["stem"])
        _set_font(r_stem)

        for opt in item.get("options", []):
            opt_p = _para(doc, left_indent=1)
            opt_run = opt_p.add_run(opt)
            _set_font(opt_run, size=10)

        doc.add_paragraph()

    _add_divider(doc)

    # ── Section B — Fill in the Blank ────────────────────────────────────────
    _heading(doc, "Section B — Fill in the Blank", level=2)

    for i, item in enumerate(sections["B"]):
        q = _question_number(sections, "B", i)
        p = _para(doc)
        r_num = p.add_run(f"{q}.  Fill in the blank with the correct word.")
        _set_font(r_num, bold=True)

        hint_p = _para(doc, left_indent=0.8)
        hint_run = hint_p.add_run(f"(synonym: {item.get('hint', '')})  —  {item['stem']}")
        _set_font(hint_run, size=10, italic=True)

        doc.add_paragraph()

    _add_divider(doc)

    # ── Section C — Error Analysis ───────────────────────────────────────────
    _heading(doc, "Section C — Error Analysis", level=2)
    instr_p = _para(doc)
    instr_run = instr_p.add_run(
        "Each sentence contains one incorrect word (shown in bold). "
        "Identify the error and write the correct word on the line."
    )
    _set_font(instr_run, size=10, italic=True)

    for i, item in enumerate(sections["C"]):
        q = _question_number(sections, "C", i)
        p = _para(doc)
        r_num = p.add_run(f"{q}.  ")
        _set_font(r_num, bold=True)
        # Parse **bold** markers and add inline
        parts = _parse_bold_markers(item["stem"])
        for text, is_bold in parts:
            run = p.add_run(text)
            _set_font(run, bold=is_bold)

        ans_p = _para(doc, left_indent=0.8)
        ans_run = ans_p.add_run("Correct word:  _______________")
        _set_font(ans_run, size=10)

        doc.add_paragraph()

    _add_divider(doc)

    # ── Section D — Real-World Application ───────────────────────────────────
    _heading(doc, "Section D — Real-World Application", level=2)
    instr_d = _para(doc)
    instr_d_run = instr_d.add_run(
        "Write 1–2 sentences in response to each prompt. "
        "Use the vocabulary word and underline it in your answer."
    )
    _set_font(instr_d_run, size=10, italic=True)

    for i, item in enumerate(sections["D"]):
        q = _question_number(sections, "D", i)
        p = _para(doc)
        r_num = p.add_run(f"{q}.  ")
        _set_font(r_num, bold=True)
        parts = _parse_bold_markers(item["stem"])
        for text, is_bold in parts:
            run = p.add_run(text)
            _set_font(run, bold=is_bold)

        for _ in range(2):
            line_p = _para(doc, left_indent=0.8)
            line_run = line_p.add_run("_______________________________________________")
            _set_font(line_run, size=11)

        doc.add_paragraph()

    _add_divider(doc)

    # ── Answer Key ───────────────────────────────────────────────────────────
    _heading(doc, "Answer Key  (Teacher Copy)", level=2)

    def _key_section(label, items, section_key):
        h = doc.add_paragraph()
        hr = h.add_run(label)
        _set_font(hr, bold=True, size=10)
        for idx, item in enumerate(items):
            q = _question_number(sections, section_key, idx)
            if section_key == "A":
                line = f"{q}.  {item['word']} — {item['answer']}"
            elif section_key == "B":
                line = f"{q}.  {item['word']} — {item['answer']}"
            elif section_key == "C":
                line = f"{q}.  {item['answer']}"
            else:
                line = f"{q}.  Answers will vary. The word '{item['word']}' must appear underlined."
            kp = _para(doc, left_indent=0.5)
            kr = kp.add_run(line)
            _set_font(kr, size=10)

    _key_section("Section A", sections["A"], "A")
    _key_section("Section B", sections["B"], "B")
    _key_section("Section C", sections["C"], "C")
    _key_section("Section D", sections["D"], "D")

    output_path = output_dir / f"quiz_{n:02d}.docx"
    doc.save(str(output_path))
    return output_path
